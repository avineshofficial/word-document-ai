from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any, Optional
import os
import uuid


class DocumentGenerator:
    def __init__(self):
        self.output_dir = "generated_docs"
        os.makedirs(self.output_dir, exist_ok=True)

    # Column widths in DXA  (S.No = 900, Title = 7200, Page No = 1260)
    COL_DXA = [900, 7200, 1260]

    # ── XML helpers ───────────────────────────────────────────────────────────

    def prevent_row_break(self, row):
        """Prevent a table row from splitting across two pages."""
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)

    def _fix_table_xml(self, table, col_dxa_list):
        """Force fixed column widths and correct XML schema order."""
        tbl   = table._tbl
        tblPr = tbl.tblPr
        total = sum(col_dxa_list)

        for e in tblPr.findall(qn("w:tblW")):
            tblPr.remove(e)

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(total))
        tblW.set(qn("w:type"), "dxa")
        tblStyle = tblPr.find(qn("w:tblStyle"))
        if tblStyle is not None:
            tblStyle.addnext(tblW)
        else:
            tblPr.insert(0, tblW)

        old_grid = tbl.find(qn("w:tblGrid"))
        if old_grid is not None:
            tbl.remove(old_grid)

        tblGrid = OxmlElement("w:tblGrid")
        for dxa in col_dxa_list:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(dxa))
            tblGrid.append(gc)
        tblPr.addnext(tblGrid)

    def _set_cell_width(self, cell, dxa: int):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for e in tcPr.findall(qn("w:tcW")):
            tcPr.remove(e)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(dxa))
        tcW.set(qn("w:type"), "dxa")
        tcPr.insert(0, tcW)

    def _set_cell_shading(self, cell, fill: str = "D3D3D3"):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for e in tcPr.findall(qn("w:shd")):
            tcPr.remove(e)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is not None:
            borders.addnext(shd)
        else:
            tcPr.append(shd)

    def _set_cell_borders(self, cell):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for e in tcPr.findall(qn("w:tcBorders")):
            tcPr.remove(e)
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            tcBorders.append(el)
        tcPr.append(tcBorders)

    def _set_cell_margins(self, cell, top=80, left=120, bottom=80, right=120):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for e in tcPr.findall(qn("w:tcMar")):
            tcPr.remove(e)
        tcMar = OxmlElement("w:tcMar")
        for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"),    str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    def _style_toc_cell(self, cell, dxa, bold=False,
                        align=WD_ALIGN_PARAGRAPH.LEFT,
                        text="", indent=None, shading=None):
        self._set_cell_width(cell, dxa)
        self._set_cell_borders(cell)
        self._set_cell_margins(cell)
        if shading:
            self._set_cell_shading(cell, shading)
        p = cell.paragraphs[0]
        p.alignment = align
        if indent:
            p.paragraph_format.left_indent = indent
        run = p.add_run(text)
        run.bold           = bold
        run.font.name      = "Times New Roman"
        run.font.size      = Pt(11)
        return run

    # ── Table of Contents ─────────────────────────────────────────────────────

    def _add_table_of_contents(self, doc, toc_headings):
        """Render the TABLE OF CONTENTS heading + bordered grid table."""
        # Section title
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.paragraph_format.space_before = Twips(200)
        toc_title.paragraph_format.space_after  = Twips(300)
        tr = toc_title.add_run("TABLE OF CONTENTS")
        tr.font.bold  = True
        tr.font.size  = Pt(14)
        tr.font.name  = "Times New Roman"

        # Table
        table = doc.add_table(rows=1, cols=3)
        table.style   = "Table Grid"
        table.autofit = False
        self._fix_table_xml(table, self.COL_DXA)

        # Header row
        hdr = table.rows[0].cells
        for cell, dxa, txt, align in [
            (hdr[0], self.COL_DXA[0], "S. No",   WD_ALIGN_PARAGRAPH.CENTER),
            (hdr[1], self.COL_DXA[1], "TITLE",    WD_ALIGN_PARAGRAPH.LEFT),
            (hdr[2], self.COL_DXA[2], "PAGE NO",  WD_ALIGN_PARAGRAPH.CENTER),
        ]:
            cell.text = ""
            self._style_toc_cell(cell, dxa, bold=True, align=align,
                                 text=txt, shading="D3D3D3")

        # Data rows
        for heading in toc_headings:
            row   = table.add_row()
            self.prevent_row_break(row)
            cells = row.cells
            sno   = str(heading.get("sno",   ""))
            text  = str(heading.get("text",  ""))
            level = heading.get("level", 1)
            is_h1 = (level == 1)

            self._style_toc_cell(cells[0], self.COL_DXA[0],
                                 bold=is_h1, align=WD_ALIGN_PARAGRAPH.CENTER, text=sno)
            self._style_toc_cell(cells[1], self.COL_DXA[1],
                                 bold=is_h1,
                                 text=text.upper() if is_h1 else text,
                                 indent=Inches(0.35) if not is_h1 else None)
            self._style_toc_cell(cells[2], self.COL_DXA[2],
                                 align=WD_ALIGN_PARAGRAPH.CENTER, text="")

        doc.add_page_break()

    # ── Document builders ─────────────────────────────────────────────────────

    def create_document(self,
                        toc_headings: List[Dict],
                        content_data: Dict,
                        student_info: Optional[Dict] = None) -> str:
        """
        Build the full project-report Word document.

        Page layout (matches SecureChain PDF):
          Page 1  — Cover
          Page 2  — Table of Contents  (page break at end)
          Page 3+ — Content body
                    • Page break BEFORE each chapter title (except the first)
                    • Chapter title centred at top of new page, 14 pt bold
                    • Sub-sections flow directly below — NO break between them
                    • Next chapter title triggers the next page break
        """
        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        # Page margins
        sec = doc.sections[0]
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)   # binding side
        sec.right_margin  = Inches(0.75)

        # ── Cover page ────────────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Inches(2)
        run = title_p.add_run("PROJECT REPORT\n\n")
        run.bold       = True
        run.font.name  = "Times New Roman"
        run.font.size  = Pt(22)

        if student_info and student_info.get("name"):
            name_p = doc.add_paragraph()
            name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nr = name_p.add_run(f"Submitted by:\n{student_info['name']}")
            nr.font.name = "Times New Roman"
            nr.font.size = Pt(16)

        doc.add_page_break()

        # ── Table of Contents ─────────────────────────────────────────────────
        self._add_table_of_contents(doc, toc_headings)

        # ── Content body ──────────────────────────────────────────────────────
        first_chapter = True

        for heading in toc_headings:
            h_text  = heading.get("text",  "")
            h_sno   = str(heading.get("sno", ""))
            h_level = heading.get("level", 1)

            if h_level == 1:
                # Page break before every chapter except the very first
                if first_chapter:
                    first_chapter = False
                else:
                    doc.add_page_break()

                # Chapter title — centred, 14 pt bold
                h_para = doc.add_paragraph()
                h_para.paragraph_format.space_before = Pt(24)
                h_para.paragraph_format.space_after  = Pt(18)
                h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = h_para.add_run(f"CHAPTER {h_sno} - {h_text.upper()}")
                run.bold      = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(14)
                # Sub-sections follow on the SAME page — no page break here
                continue

            # ── Sub-section heading (level 2+) ────────────────────────────────
            h_para = doc.add_paragraph()
            h_para.paragraph_format.space_before = Pt(18)
            h_para.paragraph_format.space_after  = Pt(10)
            h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = h_para.add_run(f"{h_sno} {h_text}")
            run.bold      = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

            # Body text
            content = content_data.get(h_text, {}).get("content", "")
            if not content.strip():
                continue   # not yet generated — skip silently

            for para_text in content.split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip())
                    p.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFIED
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.space_before      = Pt(0)
                    p.paragraph_format.space_after       = Pt(6)
                    p_run = p.runs[0] if p.runs else p.add_run()
                    p_run.font.name = "Times New Roman"
                    p_run.font.size = Pt(12)

        filename = f"Project_Report_{uuid.uuid4().hex[:6]}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def create_toc_only_document(self, toc_headings: List[Dict]) -> str:
        """
        Export a Word document that contains only the Table of Contents table.
        Called by /api/export-toc  (TocEditor.tsx 'Get TOC Docx' button).
        """
        doc = Document()
        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        sec = doc.sections[0]
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(0.75)

        self._add_table_of_contents(doc, toc_headings)

        filename = f"TOC_Export_{uuid.uuid4().hex[:6]}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath